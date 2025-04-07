# Import required modules
import os
import cv2
from PIL import Image
from ultralytics import YOLO
import pytesseract
import fitz  # PyMuPDF
import io
import joblib
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches
import cohere
import torch

import sys
import types


# Workaround for streamlit attempting to access torch.classes.__path__._path
if not isinstance(torch.classes, types.ModuleType):
    sys.modules['torch.classes'] = types.SimpleNamespace()

# === Initialize Cohere ===
cohere_api_key = "WIArC8GebX86Vj6SxyRkQVc7VcEfKTe2Dt3BFAV0"
co = cohere.Client(cohere_api_key)

# === Load Models and Encoders ===
yolo_model = YOLO(r"./best.pt")
best_model = joblib.load(r"./model_treatment_plan1.pkl")
label_encoders = joblib.load(r"./label_encoders1.pkl")
target_encoder = joblib.load(r"./target_encoders1.pkl")

pixel_to_cm = 0.07

def classify_tumor_stage(size_cm):
    if size_cm <= 3:
        return "T1"
    elif size_cm <= 5:
        return "T2"
    elif size_cm <= 7:
        return "T3"
    else:
        return "T4"

def analyze_tumor(image_path):
    results = yolo_model.predict(image_path, imgsz=640, conf=0.25)
    original_image = cv2.imread(image_path)
    height, width, _ = original_image.shape
    detection_data = []

    for box in results[0].boxes:
        x1, y1, x2, y2 = map(int, box.xyxy[0].tolist())
        box_width_px = x2 - x1
        box_height_px = y2 - y1
        area_px = box_width_px * box_height_px
        center_x = x1 + box_width_px / 2

        cv2.rectangle(original_image, (x1, y1), (x2, y2), (0, 255, 0), 3)

        width_cm = box_width_px * pixel_to_cm
        height_cm = box_height_px * pixel_to_cm
        area_cm2 = width_cm * height_cm
        size_cm = max(width_cm, height_cm)

        side = "Left Lung" if center_x < width / 2 else "Right Lung"
        confidence = box.conf[0].item()
        stage = classify_tumor_stage(size_cm)

        detection_data.append({
            "tumor_location": side,
            "tumor_size_cm": round(size_cm, 2),
            "tumor_area_cm2": round(area_cm2, 2),
            "tumor_stage": stage,
            "confidence": round(confidence * 100, 2)
        })

    annotated_path = "annotated_image.jpg"
    cv2.imwrite(annotated_path, original_image)
    return detection_data, annotated_path

def extract_biomarker_results_from_pdf(pdf_bytes):
    biomarkers = ['EGFR', 'ALK', 'KRAS', 'PD-L1', 'BRAF']
    result_dict = {}

    doc = fitz.open("pdf", pdf_bytes)
    first_page = doc.load_page(0)
    pix = first_page.get_pixmap(dpi=300)
    image_bytes = pix.tobytes("png")
    image = Image.open(io.BytesIO(image_bytes))

    text = pytesseract.image_to_string(image)

    for line in text.splitlines():
        for biomarker in biomarkers:
            if biomarker in line:
                parts = line.split()
                if len(parts) >= 2 and parts[0] == biomarker:
                    value = parts[1].capitalize()
                    if value in ["Positive", "Negative"]:
                        result_dict[biomarker] = value

    return result_dict

def prepare_model_input(tumor_features, cancer_type, chronic_conditions, biomarkers):
    return {
        "tumor_location": tumor_features["tumor_location"],
        "tumor_size_cm": tumor_features["tumor_size_cm"],
        "tumor_area_cm2": tumor_features["tumor_area_cm2"],
        "tumor_stage": tumor_features["tumor_stage"],
        "cancer_type": cancer_type,
        "chronic_condition": ", ".join(chronic_conditions),
        "EGFR": biomarkers.get("EGFR", 0),
        "ALK": biomarkers.get("ALK", 0),
        "KRAS": biomarkers.get("KRAS", 0),
        "PD-L1": biomarkers.get("PD-L1", 0),
        "BRAF": biomarkers.get("BRAF", 0),
    }

def encode_sample_input(sample_input, encoders):
    df = pd.DataFrame([sample_input])
    for column, le in encoders.items():
        if column in df.columns:
            df[column] = le.transform(df[column])
    return df

def generate_summary_with_cohere(tumor_data, biomarkers, cancer_type, chronic_condition, treatment_plan):
    prompt = f"""Generate a detailed medical report in simple, human-understandable language based on the following:
Tumor Location: {tumor_data['tumor_location']}
Tumor Size (cm): {tumor_data['tumor_size_cm']}
Tumor Area (cm²): {tumor_data['tumor_area_cm2']}
Tumor Stage: {tumor_data['tumor_stage']}
Cancer Type: {cancer_type}
Chronic Conditions: {', '.join(chronic_condition)}
Biomarkers: {', '.join([f"{k}: {v}" for k, v in biomarkers.items()])}
Recommended Treatment Plan: {treatment_plan}

Write a professional but easy-to-read summary that can be shared with the patient and medical staff. dont give patient overview
"""

    response = co.generate(
        model="command-r-plus",
        prompt=prompt,
        max_tokens=500,
        temperature=0.7,
    )
    return response.generations[0].text.strip()

def generate_word_report(tumor_data, biomarkers, cancer_type, chronic_condition, treatment_plan, image_path, summary_text):
    doc = Document()
    doc.add_heading("Lung Cancer Report", level=1)

    doc.add_heading("Generated Summary", level=2)
    doc.add_paragraph(summary_text)

    doc.add_heading("Tumor Analysis", level=2)
    for k, v in tumor_data.items():
        doc.add_paragraph(f"{k.replace('_', ' ').capitalize()}: {v}")

    doc.add_heading("Biomarkers", level=2)
    for k, v in biomarkers.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Patient Info", level=2)
    doc.add_paragraph(f"Cancer Type: {cancer_type}")
    doc.add_paragraph(f"Chronic Conditions: {', '.join(chronic_condition)}")

    doc.add_heading("Recommended Treatment Plan", level=2)
    doc.add_paragraph(treatment_plan)

    doc.add_heading("Tumor Detection Image", level=2)
    doc.add_picture(image_path, width=Inches(5))

    output_path = "Lung_Cancer_Report.docx"
    doc.save(output_path)
    return output_path

# === Streamlit UI ===
st.set_page_config(page_title="Lung Tumor Detection & Treatment", layout="centered")
st.title("🫁 Lung Tumor Detection & Treatment Plan Predictor")
st.markdown("Upload a CT scan image to analyze the tumor and generate a personalized treatment report.")

if "tumor" not in st.session_state:
    st.session_state.tumor = None
if "biomarkers" not in st.session_state:
    st.session_state.biomarkers = None
if "treatment_plan" not in st.session_state:
    st.session_state.treatment_plan = None
if "annotated_image_path" not in st.session_state:
    st.session_state.annotated_image_path = None

image_file = st.file_uploader("🖼️ Upload Lung CT Image", type=["jpg", "jpeg", "png"])

if image_file:
    image_path = "uploaded_ct.jpg"
    with open(image_path, "wb") as f:
        f.write(image_file.read())

    with st.spinner("🔍 Analyzing tumor..."):
        tumor_data, annotated_image_path = analyze_tumor(image_path)

    if tumor_data:
        tumor = tumor_data[0]
        st.session_state.tumor = tumor
        st.session_state.annotated_image_path = annotated_image_path
        st.image(annotated_image_path, caption="✅ Tumor Detected", use_column_width=True)
        st.subheader("📊 Tumor Analysis")
        st.json(tumor)

if st.session_state.tumor:
    with st.expander("🧬 View Treatment Plan", expanded=True):
        st.subheader("📑 Upload Biomarker Report (PDF)")
        pdf_file = st.file_uploader("Upload your biomarker test report", type=["pdf"])

        st.subheader("📋 Select Cancer Details")
        cancer_type = st.selectbox("Cancer Type", ["Non-Small Cell Lung Cancer", "Small Cell Lung Cancer"])
        chronic_options = ["None", "COPD", "Cardiovascular Disease (CVD)", "Kidney Disease", "Liver disease", "Autoimmune disorders"]
        chronic_condition = st.multiselect("Chronic Conditions", chronic_options, default=["None"])

        if pdf_file:
            with st.spinner("🧠 Extracting biomarkers..."):
                pdf_bytes = pdf_file.read()
                st.session_state.biomarkers = extract_biomarker_results_from_pdf(pdf_bytes)
                st.success("🧬 Biomarkers Extracted")

        if st.session_state.biomarkers:
            st.write("Extracted Biomarkers")
            st.json(st.session_state.biomarkers)

        if cancer_type and chronic_condition and st.session_state.biomarkers:
            try:
                model_input = prepare_model_input(
                    st.session_state.tumor,
                    cancer_type,
                    chronic_condition,
                    st.session_state.biomarkers
                )
                encoded_input = encode_sample_input(model_input, label_encoders)
                prediction = best_model.predict(encoded_input)
                treatment_plan = target_encoder.inverse_transform(prediction)[0]
                st.session_state.treatment_plan = treatment_plan

                st.subheader("💊 Predicted Treatment Plan")
                st.success(f"Recommended Plan: **{treatment_plan}**")

                with st.spinner("✍️ Generating detailed report with Cohere..."):
                    summary = generate_summary_with_cohere(
                        st.session_state.tumor,
                        st.session_state.biomarkers,
                        cancer_type,
                        chronic_condition,
                        treatment_plan
                    )

                report_path = generate_word_report(
                    st.session_state.tumor,
                    st.session_state.biomarkers,
                    cancer_type,
                    chronic_condition,
                    treatment_plan,
                    st.session_state.annotated_image_path,
                    summary
                )
                with open(report_path, "rb") as f:
                    st.download_button("📥 Download Word Report", f, file_name="Lung_Cancer_Report.docx")

                st.markdown("📄 **Generated Summary**")
                st.write(summary)

            except Exception as e:
                st.error(f"Prediction failed: {e}")
        else:
            st.info("⬆️ Please upload PDF and select additional inputs to proceed.")
else:
    st.info("📤 Please upload and analyze a CT scan first.")
